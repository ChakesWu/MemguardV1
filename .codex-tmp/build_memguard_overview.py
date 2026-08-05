from __future__ import annotations

import importlib.util
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Documents" / "MemGuard_產品概覽_2026-07-31.docx"
WORK = ROOT / "tmp" / "docx_build" / "memguard_overview"
WORK.mkdir(parents=True, exist_ok=True)

TABLE_HELPER = Path(
    "/Users/chakeswu/.codex/plugins/cache/openai-primary-runtime/documents/26.731.11130/"
    "skills/documents/scripts/table_geometry.py"
)
spec = importlib.util.spec_from_file_location("table_geometry", TABLE_HELPER)
table_geometry = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(table_geometry)


NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 104, 117)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
GREEN = RGBColor(28, 111, 82)
AMBER = RGBColor(145, 93, 0)
RED = RGBColor(155, 28, 28)
FONT_LATIN = "Calibri"
FONT_CJK = "Arial Unicode MS"


def set_run_font(run, size=None, bold=None, color=None, italic=None, latin=FONT_LATIN):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D0D7DE", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def create_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id, fmt, text, left, hanging, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suffix, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abstract_id))
        num.append(aid)
        numbering.append(num)

    add_abstract(20, "bullet", "•", 720, 360, "Symbol")
    add_abstract(21, "decimal", "%1.", 720, 360)
    add_num(20, 20)
    add_num(21, 21)
    return 20, 21


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_body(doc, text, bold_prefix=None, color=INK, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_callout(doc, label, text, accent="2E74B5"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_border(cell, color=accent, size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=RGBColor.from_string(accent))
    r = p.add_run(text)
    set_run_font(r, color=INK)
    table_geometry.apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=9.5, color=INK)
    table_geometry.apply_table_geometry(
        table, widths, table_width_dxa=9360, indent_dxa=120,
        cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def architecture_image(path):
    w, h = 1560, 720
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font_path = "/System/Library/Fonts/STHeiti Medium.ttc"
    font_b = ImageFont.truetype(font_path, 31)
    font_s = ImageFont.truetype(font_path, 22)
    font_xs = ImageFont.truetype(font_path, 19)
    title_font = ImageFont.truetype(font_path, 27)
    draw.text((40, 26), "MemGuard 的資料流：從記憶操作到可稽核證據", fill="#0B2545", font=title_font)

    boxes = [
        (50, 145, 285, 365, "AI Agent", "LangGraph / 多代理\n讀取或寫入記憶"),
        (355, 145, 620, 365, "MemGuard SDK", "攔截 CREATE / READ\nUPDATE / DELETE / SEARCH"),
        (690, 145, 950, 365, "事件傳輸", "非阻塞佇列\nHTTP / Stdout / File"),
        (1020, 100, 1505, 410, "Control Plane", "FastAPI ingestion\nSQLite / PostgreSQL 儲存\n衝突偵測 · Decision Trace\n稽核報告 · 租戶隔離"),
    ]
    fills = ["#EAF2F8", "#E8EEF5", "#F2F4F7", "#EAF4F0"]
    for (x1, y1, x2, y2, title, body), fill in zip(boxes, fills):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=fill, outline="#9BA8B6", width=3)
        draw.text((x1 + 22, y1 + 25), title, fill="#0B2545", font=font_b)
        draw.multiline_text((x1 + 22, y1 + 87), body, fill="#354454", font=font_s, spacing=12)

    for x1, x2, y in [(285, 355, 255), (620, 690, 255), (950, 1020, 255)]:
        draw.line((x1 + 10, y, x2 - 18, y), fill="#2E74B5", width=6)
        draw.polygon([(x2 - 18, y - 12), (x2, y), (x2 - 18, y + 12)], fill="#2E74B5")

    draw.rounded_rectangle((355, 500, 1505, 655), radius=18, fill="#FFF7E6", outline="#CDA05A", width=3)
    draw.text((385, 525), "Evidence Console / API", fill="#7A4C00", font=font_b)
    draw.text((385, 580), "時間線 · 記憶 Diff · 衝突警示 · Memory IN → Output → Memory OUT · Audit Report", fill="#574423", font=font_s)
    draw.line((1260, 410, 1260, 485), fill="#2E74B5", width=6)
    draw.polygon([(1248, 485), (1260, 503), (1272, 485)], fill="#2E74B5")
    draw.text((70, 665), "設計原則：MemGuard 觀察並記錄記憶狀態，不取代原有記憶後端，也不代理 LLM 呼叫。", fill="#66727F", font=font_xs)
    img.save(path, dpi=(180, 180))


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Quiet running header/footer for multi-page report.
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("MEMGUARD  ·  PRODUCT OVERVIEW")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("2026-07-31   ·   ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")


def main():
    diagram = WORK / "architecture.png"
    architecture_image(diagram)

    doc = Document()
    configure_document(doc)
    bullet_num, decimal_num = create_numbering(doc)

    # Cover — editorial cover pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(106)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCT OVERVIEW")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("MemGuard")
    set_run_font(r, size=32, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("AI Agent 的記憶可觀測性、安全與稽核證據層")
    set_run_font(r, size=16, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.space_after = Pt(72)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run("看見 Agent 記住了什麼、哪些記憶曾在輸出生成時可用、狀態如何演變，以及記憶是否出現衝突或可疑修改。")
    set_run_font(r, size=12, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("專案快照：MemguardV1")
    set_run_font(r, size=10, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("整理日期：2026 年 7 月 31 日")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "一頁摘要", 1)
    add_callout(
        doc,
        "一句話定位",
        "MemGuard 是 AI Agent 的『記憶黑盒飛行記錄器』：它在不替換原有記憶系統的前提下，攔截並記錄每一次記憶操作，再把這些事件組成可查詢、可比較、可稽核的證據鏈。",
    )
    add_body(doc, "你正在做的核心不是另一個向量資料庫，也不是 Prompt Log。你是在建立一層專門針對 Agent Memory 的 observability 與 governance infrastructure，讓工程、資安與合規人員可以回答四個關鍵問題：")
    for item in [
        "Agent 在某一刻讀了、寫了、更新了或刪除了哪些記憶？",
        "某個輸出生成時，有哪些持久化記憶證據可被使用？之後又寫出了什麼新狀態？",
        "同一記憶是否被不同 Agent 在短時間內互相覆寫，或出現疑似污染／指令注入？",
        "能否把技術事件轉成開發者可除錯、合規人員可審查的時間線與報告？",
    ]:
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "現在已經具備的產品形態", 2)
    add_body(doc, "目前程式庫已形成一個可展示、可本機部署的端到端產品雛形：Python SDK 負責攔截；FastAPI Control Plane 負責接收、儲存與分析；Next.js Dashboard 負責呈現；Docker Compose 串起 PostgreSQL、Keycloak、後端與前端；FinCompli 則提供具體的金融合規示範案例。")

    add_heading(doc, "產品的分層願景", 2)
    add_table(
        doc,
        ["層級", "主要使用者", "要回答的問題", "現況"],
        [
            ("Tier 1｜除錯", "AI 工程師", "是哪一段記憶造成了這次行為？", "核心能力已落地"),
            ("Tier 2｜可觀測性", "平台團隊", "記憶系統如何變化、是否穩定？", "事件、時間線、Diff、統計已具雛形"),
            ("Tier 3｜可稽核性", "風控／合規", "能否用證據解釋一次 Agent 輸出？", "Decision Trace 與模板式報告已落地"),
            ("Tier 4｜治理", "CISO／管理層", "如何把記憶當作組織風險面治理？", "屬於中長期藍圖"),
        ],
        [1300, 1700, 3900, 2460],
    )
    doc.add_page_break()

    add_heading(doc, "為什麼這個產品值得做", 1)
    add_heading(doc, "Agent 的記憶已經成為新的隱藏狀態", 2)
    add_body(doc, "傳統應用的錯誤多半可以從輸入、程式碼與資料庫重現；但具有長期記憶的 Agent 會被跨回合、跨工具、跨 Agent 的狀態影響。當答案錯了，僅看 Prompt 與 LLM 回應，往往看不到真正參與生成流程的歷史記憶、SOP、使用者偏好或工作狀態。")
    add_heading(doc, "三個高價值痛點", 2)
    for text in [
        "除錯困難：開發者知道輸出錯了，卻不知道 Agent 當時讀到哪個舊狀態、是否讀到衝突版本。",
        "安全風險：惡意內容一旦被寫入長期記憶，可能在未來多次被重新讀取，形成延遲生效的 memory poisoning。",
        "合規缺口：高風險決策需要可追溯證據，但普通日誌沒有以記憶實體、版本、讀寫關係與租戶界線組織。",
    ]:
        add_list_item(doc, text, bullet_num, bold_prefix=text.split("：")[0] + "：")
    add_callout(doc, "產品洞察", "在微服務世界，Distributed Tracing 讓請求路徑可見；MemGuard 的主張是：在 Agent 世界，Memory Tracing 會成為同樣基礎的能力。", accent="1F4D78")

    add_heading(doc, "它刻意不做什麼", 2)
    for text in [
        "不取代 Mem0、Redis、向量資料庫或 LangGraph Checkpointer 等既有記憶後端。",
        "不把所有 LLM 請求代理到自己的服務；目前重點是記憶操作與證據關聯。",
        "不宣稱『記錄到的記憶』等同模型內部真正的因果解釋；它提供的是可驗證的 evidence lineage，而不是讀心術。",
    ]:
        add_list_item(doc, text, bullet_num)
    doc.add_page_break()

    add_heading(doc, "系統如何運作", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run()
    r.add_picture(str(diagram), width=Inches(6.42))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("圖 1｜MemGuard 端到端資料流")
    set_run_font(r, size=9, italic=True, color=MUTED)

    add_heading(doc, "核心流程", 2)
    steps = [
        "Agent 透過既有記憶機制進行 CREATE、READ、UPDATE、DELETE、QUERY 或 SEARCH。",
        "MemGuard SDK／框架 Adapter 在操作邊界產生 MemoryEvent；事件包含 Agent、Session、Namespace、記憶鍵、類型、時間、Hash 與上下文。",
        "Transport 以背景佇列非同步送出事件。即使觀測後端不可用，設計目標也是不阻塞 Agent 主流程。",
        "Control Plane 驗證並持久化事件，同時建立 Decision Trace、衝突分析、統計與稽核資料。",
        "Evidence Console 把原始事件轉為時間線、狀態差異、輸出前證據、輸出後寫入與報告。",
    ]
    for step in steps:
        add_list_item(doc, step, decimal_num)

    add_heading(doc, "兩個最重要的資料單位", 2)
    add_table(
        doc,
        ["資料單位", "代表什麼", "核心欄位／關聯"],
        [
            ("MemoryEvent", "一次原子記憶操作", "event_id、agent/session/namespace、operation、memory_key/type、before/after、content_hash、timestamp、context"),
            ("DecisionTrace", "把一次 Agent 輸出前後的證據串起來", "input_event_ids → prompt/output hash 與摘要 → output_event_ids，另可附 evidence ranking / influence score"),
        ],
        [1700, 2500, 5160],
    )
    doc.add_page_break()

    add_heading(doc, "目前已實作的產品能力", 1)
    sections = [
        ("1. SDK 與整合", [
            "可安裝的 Python SDK，核心包含 MemoryEvent、Interceptor、DecisionTrace 與 Influence 計算。",
            "LangGraph Checkpointer Adapter 可包裝既有 checkpointer，攔截 get / put 類型操作。",
            "提供 Stdout、HTTP、File 與 Null Transport；HTTP Transport 支援背景佇列、批次、重試與 flush。",
            "事件傳送採 best-effort／fire-and-forget 思路，觀測失敗不應中斷業務 Agent。",
        ]),
        ("2. Control Plane 與資料", [
            "FastAPI 提供事件 ingestion、記憶讀寫／查詢、時間線、Decision Trace、影響歷史、衝突分析、Session、統計與稽核 API。",
            "本機預設可使用 SQLite；Pilot／Compose 路徑支援 PostgreSQL 16。",
            "Keycloak OIDC Bearer Token、tenant claim 與 API tenant enforcement 已進入主架構，用來阻止跨租戶讀取。",
        ]),
        ("3. 分析與呈現", [
            "Memory Timeline：依時間查看讀寫事件並按 operation / agent 篩選。",
            "Memory Diff：顯示 before / after 以及欄位變動。",
            "Conflict Detection：辨識不同 Agent 在短時間內修改同一記憶鍵的情況。",
            "Decision Trace：呈現 Memory IN → Agent Output → Memory OUT，並保留缺失證據警示。",
            "Audit Report：以模板把技術事件整理為 compliance、debug 或 business 風格摘要。",
        ]),
    ]
    for title, items in sections:
        add_heading(doc, title, 2)
        for item in items:
            add_list_item(doc, item, bullet_num)

    add_callout(doc, "隱私與證據", "SDK 的資料模型以 content hash 為核心追蹤值；原文內容可以按整合與部署策略保留或最小化。Hash 能證明內容版本是否改變，但本身不是資料防洩漏的完整替代品。", accent="1C6F52")
    doc.add_page_break()

    add_heading(doc, "示範場景：FinCompli 金融合規 Agent", 1)
    add_body(doc, "專案內的 FinCompli baseline 是 MemGuard 的垂直示範：多個 Agent 共同分析一宗疑似拆分交易（structuring）案例。MemGuard 的價值不是替合規 Agent 下判斷，而是把判斷過程中的記憶證據留下來。")

    add_heading(doc, "示範故事", 2)
    for step in [
        "Fraud Detection Agent 讀取客戶／交易歷史，產生 fraud analysis。",
        "Case History Agent 搜尋過去 SAR 案例，取得相似案例與相似度。",
        "Compliance Research Agent 查閱規例／SOP 記憶。",
        "Report Generation Agent 綜合工作記憶，產出 FILE SAR 或其他處置建議。",
        "MemGuard 把所有讀取、搜尋、工作記憶寫入與最終輸出串成可審查鏈路。",
    ]:
        add_list_item(doc, step, decimal_num)

    add_heading(doc, "同一份證據，服務三種角色", 2)
    add_table(
        doc,
        ["角色", "他看到的價值", "典型問題"],
        [
            ("AI 工程師", "事件時間線、狀態 Diff、缺失鏈路", "哪個 checkpoint 或記憶版本讓結果偏掉？"),
            ("平台／資安", "衝突、租戶隔離、異常寫入與操作量", "是否有跨 Agent 覆寫、污染或資料邊界問題？"),
            ("合規／業務", "輸入證據、輸出摘要、後續寫入與 Audit Report", "這個建議有什麼持久化證據可供人員覆核？"),
        ],
        [1700, 3600, 4060],
    )

    add_callout(doc, "對外講法", "Without MemGuard：『AI 標記了這筆交易，但我們很難重建它當時看到的記憶。』 With MemGuard：『我們可以逐項查看輸出生成時已持久化的案例、規例與工作狀態，以及輸出後寫入了什麼。』", accent="2E74B5")
    doc.add_page_break()

    add_heading(doc, "目前成熟度、限制與下一步", 1)
    add_heading(doc, "可合理宣稱的現況", 2)
    add_table(
        doc,
        ["面向", "現在可以說", "暫時不要過度承諾"],
        [
            ("產品", "可本機展示的端到端 Memory Observability MVP", "已是完整企業治理平台"),
            ("整合", "LangGraph Adapter 與通用 Interceptor／Transport 基礎", "已支援所有主流 Agent／Memory Framework"),
            ("證據", "可重建持久化事件與輸出前後 lineage", "能證明模型內部真正的因果推理"),
            ("安全", "OIDC、租戶隔離、衝突與簡單可疑內容規則", "已具完整 RBAC、策略引擎、不可竄改帳本與企業 SIEM 整合"),
            ("稽核", "可產生模板式 session report", "已完成法規級自動合規判定"),
            ("部署", "Docker Compose：Keycloak + PostgreSQL + FastAPI + Next.js", "已完成多區域、高可用、雲端托管與 SLA"),
        ],
        [1450, 3850, 4060],
    )

    add_heading(doc, "最重要的工程限制", 2)
    for item in [
        "目前主 SDK Adapter 以 LangGraph 為核心；Mem0、AutoGen、CrewAI、Zep 等仍屬擴展方向。",
        "Influence score／evidence ranking 是解釋性排序訊號，不應包裝成已被科學驗證的模型因果貢獻值。",
        "衝突偵測目前主要依記憶鍵、Agent 與時間窗規則；語義衝突、版本合併與自動修復仍可深化。",
        "稽核摘要以模板生成為主；LLM 增強敘事在程式中保留接口，但不是完整交付能力。",
        "這次環境缺少 pytest 套件，因此本文對功能的判斷以程式碼、Docker 配置與既有測試檔為依據，未重新跑完整測試套件。",
    ]:
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "建議的下一階段", 2)
    for item in [
        "產品聚焦：先把『Output → Evidence → Memory Writes』調查工作流做成唯一主故事。",
        "證據可信度：加入不可竄改事件鏈、簽章、保留策略與 evidence completeness 指標。",
        "框架覆蓋：優先補 Mem0，再依客戶選擇 AutoGen／CrewAI／Zep。",
        "治理能力：增加 policy engine、RBAC、告警、人工覆核與 SIEM／ticketing 整合。",
        "企業部署：完成 migration、備份還原、壓力測試、可觀測性、自動化 CI 與雲端拓撲。",
    ]:
        add_list_item(doc, item, decimal_num)
    doc.add_page_break()

    add_heading(doc, "技術版圖與快速理解", 1)
    add_heading(doc, "主要技術組件", 2)
    add_table(
        doc,
        ["區塊", "技術／目錄", "責任"],
        [
            ("SDK", "sdk/memguard", "事件模型、攔截、LangGraph Adapter、Transport、Influence"),
            ("Backend", "backend/app · FastAPI", "ingestion、查詢、分析、稽核、OIDC tenant enforcement"),
            ("Database", "SQLite / PostgreSQL", "MemoryEvent、DecisionTrace 與 migration"),
            ("Identity", "Keycloak", "OIDC 登入、JWT、tenant claim"),
            ("Frontend", "frontend · Next.js", "Evidence Console、Timeline、Diff、Conflict、Audit"),
            ("Demo", "demo_simple.py / demo_with_dashboard.py / fincompli-baseline", "終端展示、Dashboard 展示、金融合規案例"),
            ("Deployment", "docker-compose.yml", "本機／Pilot 四服務編排"),
        ],
        [1550, 3150, 4660],
    )

    add_heading(doc, "如果要向別人介紹這個專案", 2)
    add_callout(doc, "30 秒版本", "MemGuard 是 AI Agent 的記憶可觀測性與安全層。它記錄 Agent 對長期與工作記憶的每一次讀寫，並把這些事件串成輸出前後的證據鏈。這讓工程師能除錯、資安團隊能發現衝突或污染、合規人員能覆核高風險 Agent 輸出。", accent="0B2545")
    add_callout(doc, "投資／產品版本", "當 Agent 開始依賴跨回合記憶，記憶就變成新的企業風險面。MemGuard 想成為這一層的 Datadog + audit trail：先從除錯與 evidence console 切入，再走向政策、權限、保留與治理。", accent="7A5A00")

    add_heading(doc, "一句話結論", 2)
    add_body(doc, "你正在把『AI 為什麼這樣做』從模糊的模型說法，轉成可查驗的記憶事件、版本差異與證據鏈。這個切入點既有立即的工程價值，也能延伸到安全與合規治理。", bold_prefix="你正在把『AI 為什麼這樣做』")

    add_heading(doc, "依據範圍", 2)
    add_body(doc, "本文件依據 2026-07-31 工作區快照整理，主要參考 README、SDK／Backend／Frontend 程式碼、Docker Compose、技術設計、產品規格、測試檔與 Dashboard redesign spec。既有較早期進度文件互相存在狀態差異，因此本文優先採用目前程式碼中可直接辨識的能力。", color=MUTED, after=0)

    doc.core_properties.title = "MemGuard 產品概覽"
    doc.core_properties.subject = "AI Agent 記憶可觀測性、安全與稽核證據層"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "MemGuard, AI Agent, Memory Observability, Security, Audit"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
